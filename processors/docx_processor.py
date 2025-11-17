# DOCX文件敏感数据识别处理器
# 通用技术文档排除，主要对图片/表格/文本分类型处理

from docx import Document
from core import idcard, bankcard, name, address, phone, email, ip

# 敏感类型函数注册表，和csv一致
SENSITIVE_TYPE_FUNCS = {
    '身份证': idcard.is_idcard,
    '银行卡': bankcard.is_bankcard,
    '姓名': name.is_name,
    '地址': address.is_address,
    '手机号': phone.is_phone,
    '邮箱': email.is_email,
    'IP': ip.is_ip,
}
SENSITIVE_MASK_FUNCS = {
    '身份证': idcard.mask_idcard,
    '银行卡': bankcard.mask_bankcard,
    '姓名': name.mask_name,
    '地址': address.mask_address,
    '手机号': phone.mask_phone,
    '邮箱': email.mask_email,
    'IP': ip.mask_ip,
}
class DOCXProcessor:
    """
    DOCX文件敏感数据识别与脱敏：
    - 遍历所有段落、表格细胞，逐条判别敏感类型
    - 支持批量按原文做脱敏，图片处理有待后续补充
    """
    def __init__(self):
        pass
    def detect_text(self, file_path: str):
        """
        遍历docx所有段落与所有表格cell，一旦命中敏感类型即记录。
        返回：段落和表格各自的敏感结构化列表
        """
        doc = Document(file_path)
        para_results = []
        # 段落遍历
        for para in doc.paragraphs:
            for stype, func in SENSITIVE_TYPE_FUNCS.items():
                if func(para.text):  # 文本只要命中即标签所属类型
                    para_results.append({'type': stype, 'content': para.text})
                    break
        table_results = []
        # 所有表格遍历所有cell
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for stype, func in SENSITIVE_TYPE_FUNCS.items():
                        if func(cell.text):
                            table_results.append({'type': stype, 'content': cell.text})
                            break
        return {'paragraph_sensitive': para_results, 'table_sensitive': table_results}
    def mask_text(self, file_path: str):
        """
        在docx原文上批量脱敏
        - 段落和表格内容会被就地修改（如有敏感立即变化），便于导出或用于预览
        """
        doc = Document(file_path)
        for para in doc.paragraphs:
            for stype, func in SENSITIVE_TYPE_FUNCS.items():
                if func(para.text):
                    mask_func = SENSITIVE_MASK_FUNCS[stype]
                    para.text = mask_func(para.text)  # 注意：对docx操作脱敏直接改文本
                    break
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for stype, func in SENSITIVE_TYPE_FUNCS.items():
                        if func(cell.text):
                            mask_func = SENSITIVE_MASK_FUNCS[stype]
                            cell.text = mask_func(cell.text)
                            break
        # 图片分支处理可按实际业务，在此补充图片敏感自动识别和标注
        # for shape in doc.inline_shapes: ...
        return doc  # 调用save可直接导出，结构化用解析结果
