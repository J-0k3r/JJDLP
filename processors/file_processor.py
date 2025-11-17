import os
import json
import csv
import io
from typing import List, Dict, Any, Optional
from abc import ABC, abstractmethod
import PyPDF2
from docx import Document
try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False
from PIL import Image
from core.detector import SensitiveDataDetector, SensitiveInfo
# 图片处理使用 idcard_ocr API

class FileProcessor(ABC):
    """文件处理器基类"""
    
    def __init__(self, detector: Optional[SensitiveDataDetector] = None):
        # 允许外部注入共享的检测器（带自定义规则），否则创建本地实例
        self.detector = detector if detector is not None else SensitiveDataDetector()
    
    @abstractmethod
    def extract_text(self, file_path: str) -> str:
        """提取文件中的文本内容"""
        pass
    
    @abstractmethod
    def get_supported_extensions(self) -> List[str]:
        """获取支持的文件扩展名"""
        pass
    
    def process_file(self, file_path: str) -> Dict[str, Any]:
        """处理文件并返回敏感信息检测结果"""
        try:
            text = self.extract_text(file_path)
            sensitive_items = self.detector.detect_sensitive_data(text)
            
            return {
                'success': True,
                'text': text,
                'sensitive_items': sensitive_items,
                'file_path': file_path,
                'file_size': os.path.getsize(file_path)
            }
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'file_path': file_path
            }

class TextProcessor(FileProcessor):
    """文本文件处理器"""
    
    def extract_text(self, file_path: str) -> str:
        """提取文本文件内容"""
        encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        
        # 如果所有编码都失败，使用二进制模式读取
        with open(file_path, 'rb') as f:
            return f.read().decode('utf-8', errors='ignore')
    
    def get_supported_extensions(self) -> List[str]:
        return ['txt', 'log', 'md', 'xml', 'html', 'htm']

class PDFProcessor(FileProcessor):
    """PDF文件处理器"""
    
    def _extract_images_from_pdf(self, file_path: str) -> List[tuple]:
        """
        从PDF中提取所有图片，返回图片数据列表。
        返回: [(image_data, page_num), ...]
        """
        images = []
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    
                    # 提取页面中的图片对象
                    try:
                        # 获取页面资源
                        if '/Resources' in page:
                            resources = page['/Resources']
                            if '/XObject' in resources:
                                xobject = resources['/XObject']
                                xobject_obj = xobject.get_object() if hasattr(xobject, 'get_object') else xobject
                                
                                # 遍历所有XObject对象
                                for obj_name in xobject_obj:
                                    obj = xobject_obj[obj_name]
                                    # 处理间接引用
                                    if hasattr(obj, 'get_object'):
                                        obj = obj.get_object()
                                    
                                    # 检查是否为图片
                                    if '/Subtype' in obj and obj['/Subtype'] == '/Image':
                                        try:
                                            # 获取图片数据
                                            image_data = None
                                            if hasattr(obj, 'get_data'):
                                                image_data = obj.get_data()
                                            elif '/Filter' in obj:
                                                # 尝试读取原始数据
                                                if hasattr(obj, '_data'):
                                                    image_data = obj._data
                                                elif '/Length' in obj:
                                                    # 使用间接方式读取
                                                    try:
                                                        stream_obj = obj if obj.get('/Type') == '/XObject' else None
                                                        if stream_obj and hasattr(stream_obj, 'get_data'):
                                                            image_data = stream_obj.get_data()
                                                    except:
                                                        pass
                                            
                                            if image_data:
                                                images.append((image_data, page_num + 1))
                                        except Exception as e:
                                            # 跳过无法提取的图片
                                            import warnings
                                            warnings.warn(f"提取PDF第{page_num+1}页图片失败: {e}")
                                            continue
                    except Exception as e:
                        import warnings
                        warnings.warn(f"提取PDF第{page_num+1}页资源失败: {e}")
                        continue
        except Exception as e:
            import warnings
            warnings.warn(f"PDF图片提取失败: {e}")
        
        return images
    
    def extract_text(self, file_path: str) -> str:
        """提取PDF文件中的文本和图片信息"""
        text = ""
        image_results = []
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # 提取文本
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            # 提取并识别图片
            images = self._extract_images_from_pdf(file_path)
            if images:
                from processors.image_processor import classify_image
                import tempfile
                
                idcard_count = 0
                bankcard_count = 0
                
                for idx, (image_data, page_num) in enumerate(images):
                    # 保存临时图片文件进行识别
                    try:
                        with tempfile.NamedTemporaryFile(delete=False, suffix='.jpg') as tmp_file:
                            tmp_file.write(image_data)
                            tmp_path = tmp_file.name
                        
                        # 识别图片类型
                        image_type = classify_image(tmp_path)
                        
                        # 统计数量
                        if "包含身份证图片" in image_type:
                            idcard_count += 1
                        elif "包含银行卡图片" in image_type:
                            bankcard_count += 1
                        
                        # 清理临时文件
                        try:
                            os.unlink(tmp_path)
                        except:
                            pass
                    except Exception as e:
                        import warnings
                        warnings.warn(f"识别PDF第{page_num}页图片失败: {e}")
                        continue
                
                # 生成图片识别结果文本
                if idcard_count > 0 or bankcard_count > 0:
                    if idcard_count > 0 and bankcard_count > 0:
                        image_results.append(f"包含{idcard_count}张身份证图片，包含{bankcard_count}张银行卡图片")
                    elif idcard_count > 0:
                        image_results.append(f"包含{idcard_count}张身份证图片")
                    elif bankcard_count > 0:
                        image_results.append(f"包含{bankcard_count}张银行卡图片")
        
        except Exception as e:
            import warnings
            warnings.warn(f"PyPDF2提取失败: {e}")
        
        # 组合文本和图片识别结果
        result_parts = []
        if text.strip():
            result_parts.append(text)
        if image_results:
            result_parts.extend(image_results)
        
        return "\n".join(result_parts) if result_parts else ""
    
    def get_supported_extensions(self) -> List[str]:
        return ['pdf']

class DOCXProcessor(FileProcessor):
    """DOCX文件处理器"""
    
    def _extract_images_from_docx(self, file_path: str) -> List[bytes]:
        """
        从DOCX中提取所有图片，返回图片数据列表。
        返回: [image_data, ...]
        """
        images = []
        try:
            import zipfile
            from xml.etree import ElementTree as ET
            
            # DOCX文件实际上是一个ZIP文件
            with zipfile.ZipFile(file_path, 'r') as docx_zip:
                # 查找媒体文件（通常在word/media/目录下）
                media_files = [name for name in docx_zip.namelist() if name.startswith('word/media/')]
                
                for media_file in media_files:
                    # 只处理图片文件
                    if any(media_file.lower().endswith(ext) for ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']):
                        try:
                            image_data = docx_zip.read(media_file)
                            images.append(image_data)
                        except Exception as e:
                            import warnings
                            warnings.warn(f"提取DOCX图片失败 {media_file}: {e}")
                            continue
        except Exception as e:
            import warnings
            warnings.warn(f"DOCX图片提取失败: {e}")
        
        return images
    
    def extract_text(self, file_path: str) -> str:
        """提取DOCX文件中的文本和图片信息"""
        text = ""
        image_results = []
        
        try:
            doc = Document(file_path)
            
            # 提取文本
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # 提取表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            # 提取并识别图片
            images = self._extract_images_from_docx(file_path)
            if images:
                from processors.image_processor import classify_image
                import tempfile
                
                idcard_count = 0
                bankcard_count = 0
                
                for idx, image_data in enumerate(images):
                    # 保存临时图片文件进行识别
                    try:
                        # 根据图片数据的前几个字节判断格式
                        suffix = '.jpg'
                        if image_data[:8] == b'\x89PNG\r\n\x1a\n':
                            suffix = '.png'
                        elif image_data[:2] == b'BM':
                            suffix = '.bmp'
                        elif image_data[:6] in [b'GIF87a', b'GIF89a']:
                            suffix = '.gif'
                        
                        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
                            tmp_file.write(image_data)
                            tmp_path = tmp_file.name
                        
                        # 识别图片类型
                        image_type = classify_image(tmp_path)
                        
                        # 统计数量
                        if "包含身份证图片" in image_type:
                            idcard_count += 1
                        elif "包含银行卡图片" in image_type:
                            bankcard_count += 1
                        
                        # 清理临时文件
                        try:
                            os.unlink(tmp_path)
                        except:
                            pass
                    except Exception as e:
                        import warnings
                        warnings.warn(f"识别DOCX图片失败: {e}")
                        continue
                
                # 生成图片识别结果文本
                if idcard_count > 0 or bankcard_count > 0:
                    if idcard_count > 0 and bankcard_count > 0:
                        image_results.append(f"包含{idcard_count}张身份证图片，包含{bankcard_count}张银行卡图片")
                    elif idcard_count > 0:
                        image_results.append(f"包含{idcard_count}张身份证图片")
                    elif bankcard_count > 0:
                        image_results.append(f"包含{bankcard_count}张银行卡图片")
        
        except Exception as e:
            raise Exception(f"DOCX文件处理失败: {e}")
        
        # 组合文本和图片识别结果
        result_parts = []
        if text.strip():
            result_parts.append(text)
        if image_results:
            result_parts.extend(image_results)
        
        return "\n".join(result_parts) if result_parts else ""
    
    def get_supported_extensions(self) -> List[str]:
        return ['docx', 'doc']

class JSONProcessor(FileProcessor):
    """JSON文件处理器"""
    
    def extract_text(self, file_path: str) -> str:
        """提取JSON文件中的文本内容"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # 递归提取JSON中的所有文本值
            def extract_text_from_dict(obj):
                text_parts = []
                
                if isinstance(obj, dict):
                    for value in obj.values():
                        text_parts.extend(extract_text_from_dict(value))
                elif isinstance(obj, list):
                    for item in obj:
                        text_parts.extend(extract_text_from_dict(item))
                elif isinstance(obj, str):
                    text_parts.append(obj)
                
                return text_parts
            
            text_parts = extract_text_from_dict(data)
            return " ".join(text_parts)
        
        except Exception as e:
            raise Exception(f"JSON文件处理失败: {e}")
    
    def get_supported_extensions(self) -> List[str]:
        return ['json']

class CSVProcessor(FileProcessor):
    """CSV文件处理器"""
    
    def extract_text(self, file_path: str) -> str:
        """提取CSV文件中的文本内容"""
        try:
            if PANDAS_AVAILABLE:
                # 使用pandas处理
                encodings = ['utf-8', 'gbk', 'gb2312']
                
                for encoding in encodings:
                    try:
                        df = pd.read_csv(file_path, encoding=encoding)
                        break
                    except UnicodeDecodeError:
                        continue
                else:
                    df = pd.read_csv(file_path, encoding='utf-8', errors='ignore')
                
                # 将所有数据转换为文本
                text_parts = []
                for column in df.columns:
                    text_parts.append(str(column))
                
                for _, row in df.iterrows():
                    for value in row.values:
                        if pd.notna(value):
                            text_parts.append(str(value))
                
                return " ".join(text_parts)
            else:
                # 使用标准库处理CSV
                encodings = ['utf-8', 'gbk', 'gb2312']
                text_parts = []
                
                for encoding in encodings:
                    try:
                        with open(file_path, 'r', encoding=encoding, newline='') as csvfile:
                            reader = csv.reader(csvfile)
                            for row in reader:
                                for cell in row:
                                    if cell.strip():  # 忽略空单元格
                                        text_parts.append(cell.strip())
                        break
                    except UnicodeDecodeError:
                        continue
                else:
                    # 如果所有编码都失败，使用错误处理
                    with open(file_path, 'r', encoding='utf-8', errors='ignore', newline='') as csvfile:
                        reader = csv.reader(csvfile)
                        for row in reader:
                            for cell in row:
                                if cell.strip():
                                    text_parts.append(cell.strip())
                
                return " ".join(text_parts)
        
        except Exception as e:
            raise Exception(f"CSV文件处理失败: {e}")
    
    def get_supported_extensions(self) -> List[str]:
        return ['csv']

class ImageProcessor(FileProcessor):
    """图片文件处理器（使用 idcard_ocr API 进行图片类型识别）"""
    
    def __init__(self, detector: Optional[SensitiveDataDetector] = None):
        super().__init__(detector)
    
    def extract_text(self, file_path: str) -> str:
        """
        通过调用 idcard_ocr API 识别图片类型。
        返回：包含身份证图片 / 包含银行卡图片 / 不包含银行卡和身份证图片
        """
        try:
            from processors.image_processor import classify_image
            image_type = classify_image(file_path)
            # 直接返回识别结果
            return image_type
        except Exception as e:
            # 捕获所有异常，返回默认值
            import warnings
            warnings.warn(f"图片类型识别失败: {str(e)}")
            return "不包含银行卡和身份证图片"
    
    def get_supported_extensions(self) -> List[str]:
        return ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff']

class FileProcessorFactory:
    """文件处理器工厂"""
    
    def __init__(self, detector: Optional[SensitiveDataDetector] = None):
        # 使用共享检测器，使自定义规则在文件扫描中生效
        self._detector = detector if detector is not None else SensitiveDataDetector()
        self.processors = {
            'text': TextProcessor(self._detector),
            'pdf': PDFProcessor(self._detector),
            'docx': DOCXProcessor(self._detector),
            'json': JSONProcessor(self._detector),
            'csv': CSVProcessor(self._detector),
            'image': ImageProcessor(self._detector)
        }
    
    def get_processor(self, file_path: str) -> Optional[FileProcessor]:
        """根据文件扩展名获取对应的处理器"""
        file_extension = file_path.lower().split('.')[-1]
        
        for processor_type, processor in self.processors.items():
            if file_extension in processor.get_supported_extensions():
                return processor
        
        return None
    
    def process_file(self, file_path: str) -> Dict[str, Any]:
        """处理文件"""
        processor = self.get_processor(file_path)
        
        if processor is None:
            return {
                'success': False,
                'error': f'不支持的文件格式: {file_path}',
                'file_path': file_path
            }
        
        return processor.process_file(file_path)

# 使用示例
if __name__ == "__main__":
    factory = FileProcessorFactory()
    
    # 测试不同文件类型
    test_files = [
        'test.txt',
        'test.pdf',
        'test.docx',
        'test.json',
        'test.csv',
        'test.jpg'
    ]
    
    for file_path in test_files:
        if os.path.exists(file_path):
            result = factory.process_file(file_path)
            print(f"处理 {file_path}: {result['success']}")
            if result['success']:
                print(f"发现 {len(result['sensitive_items'])} 个敏感信息")
            else:
                print(f"错误: {result['error']}")
